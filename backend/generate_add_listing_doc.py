from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime

def main():
    doc = Document()
    
    # Title
    doc.add_heading('Test Report', level=1)
    
    # Create table with 6 columns
    table = doc.add_table(rows=0, cols=6)
    table.style = 'Table Grid'
    
    def add_merged_row(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
        row = table.add_row()
        cell = row.cells[0]
        for i in range(1, 6):
            cell.merge(row.cells[i])
        cell.text = text
        if bold:
            cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = align
        return cell

    def add_split_row(text1, text2, bold=False):
        row = table.add_row()
        # merge 0,1,2
        cell1 = row.cells[0]
        cell1.merge(row.cells[1])
        cell1.merge(row.cells[2])
        cell1.text = text1
        if bold and cell1.paragraphs[0].runs:
            cell1.paragraphs[0].runs[0].bold = True
        
        # merge 3,4,5
        cell2 = row.cells[3]
        cell2.merge(row.cells[4])
        cell2.merge(row.cells[5])
        cell2.text = text2
        if bold and cell2.paragraphs[0].runs:
            cell2.paragraphs[0].runs[0].bold = True
        return row
    
    # Row 0
    add_merged_row('Test Case 2', bold=True)
    # Row 1
    add_merged_row('Project Name: Bidverse', bold=True)
    # Row 2
    add_merged_row('Add Listing Test Case', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    current_date = datetime.datetime.now().strftime("%d/%m/%Y")
    
    # Row 3
    add_split_row('Test Case ID: Test_2', 'Test Designed By: Jewel Treasa Raphel', bold=True)
    add_split_row('Test Priority(Low/Medium/High): High', f'Test Designed Date: {current_date}', bold=True)
    add_split_row('Module Name: Listings Module', 'Test Executed By: Sona Maria Sebastian', bold=True)
    add_split_row('Test Title: Verify user can add a new listing', f'Test Execution Date: {current_date}', bold=True)
    
    # Row 7 (Description)
    desc_row = add_split_row('Description: To verify that a logged in user can successfully navigate to the Add Listed section and add a new listing with valid details.', '', bold=True)
    
    # Row 8
    add_merged_row('Pre-Condition: User is logged in to the application', bold=True)
    
    # Headers
    hdr_row = table.add_row()
    headers = ["Step", "Test Step", "Test Data", "Expected Result", "Actual Result", "Status(Pass/Fail)"]
    for i, h in enumerate(headers):
        hdr_row.cells[i].text = h
        hdr_row.cells[i].paragraphs[0].runs[0].bold = True
        
    steps_data = [
        (
            "1",
            "Login as seller",
            "Email: alan@gmail.com\nPassword: alan@123",
            "User successfully logs in and is on the home page",
            "Login successful",
            "Pass"
        ),
        (
            "2",
            "Navigate to Dashboard",
            "N/A",
            "Dashboard page is loaded",
            "Dashboard loaded",
            "Pass"
        ),
        (
            "3",
            "Click on 'Add Listed' section from Dashboard sidebar",
            "N/A",
            "Add Listing form is loaded and visible",
            "Add Listing form page is displayed",
            "Pass"
        ),
        (
            "4",
            "Fill in commodity name",
            "Organic Wheat",
            "Commodity field contains 'Organic Wheat'",
            "Filled successfully",
            "Pass"
        ),
        (
            "5",
            "Fill in quantity and select unit",
            "Quantity: 100\nUnit: kg",
            "Quantity is 100, unit is set to kg",
            "Filled successfully",
            "Pass"
        ),
        (
            "6",
            "Fill in base price",
            "45.50",
            "Base price field contains '45.50'",
            "Filled successfully",
            "Pass"
        ),
        (
            "7",
            "Fill in description",
            "High quality organic wheat freshly harvested from the farm.",
            "Description field contains text",
            "Filled successfully",
            "Pass"
        ),
        (
            "8",
            "Select date (tomorrow) and Morning/Evening session",
            "Date: Tomorrow's date\nMorning: Unchecked\nEvening: Checked",
            "Date is selected, Evening session is checked",
            "Selected successfully",
            "Pass"
        ),
        (
            "9",
            "Click Submit button to publish the listing",
            "N/A",
            "Listing is submitted successfully",
            "Form submitted",
            "Pass"
        ),
        (
            "10",
            "Verify redirection to My Listings",
            "URL contains 'section=listings'",
            "Redirected to listings page successfully",
            "Redirected successfully",
            "Pass"
        )
    ]
    
    for row_data in steps_data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text
            
    # Post-condition
    add_merged_row('Post-Condition: New listing is created and visible in the My Listings section.', bold=True)
    
    # Adjust widths
    for cell in table.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in table.columns[1].cells:
        cell.width = Inches(1.5)
    for cell in table.columns[2].cells:
        cell.width = Inches(1.2)
    for cell in table.columns[3].cells:
        cell.width = Inches(1.5)
    for cell in table.columns[4].cells:
        cell.width = Inches(1.5)
    for cell in table.columns[5].cells:
        cell.width = Inches(0.8)
        
    doc.save('Add_Listing_Test_Case.docx')
    print("Document successfully created!")

if __name__ == "__main__":
    main()
