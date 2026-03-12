from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 12, "color": "#000000", "val": "single"},
        start={"sz": 12, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def main():
    doc = Document()
    
    # Title
    heading = doc.add_heading('Test Report', level=1)
    
    # Create table with 6 columns
    table = doc.add_table(rows=0, cols=6)
    table.style = 'Table Grid'
    
    def add_merged_row(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
        row = table.add_row()
        cell = row.cells[0]
        for i in range(1, 6):
            cell.merge(row.cells[i])
        cell.text = text
        if bold:
            cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = align
        return cell

    def add_split_row(text1, text2, bold=False):
        row = table.add_row()
        # merge 0,1,2
        cell1 = row.cells[0]
        cell1.merge(row.cells[1])
        cell1.merge(row.cells[2])
        cell1.text = text1
        if bold and cell1.paragraphs[0].runs:
            cell1.paragraphs[0].runs[0].bold = True
        
        # merge 3,4,5
        cell2 = row.cells[3]
        cell2.merge(row.cells[4])
        cell2.merge(row.cells[5])
        cell2.text = text2
        if bold and cell2.paragraphs[0].runs:
            cell2.paragraphs[0].runs[0].bold = True
        return row
    
    # Row 0
    add_merged_row('Test Case 1', bold=True)
    # Row 1
    add_merged_row('Project Name: Bidverse', bold=True)
    # Row 2
    add_merged_row('Login Test Case', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Row 3
    add_split_row('Test Case ID: Test_1', 'Test Designed By: Jewel Treasa Raphel', bold=True)
    add_split_row('Test Priority(Low/Medium/High): High', 'Test Designed Date: 06/02/2026', bold=True)
    add_split_row('Module Name: Login Module', 'Test Executed By: Sona Maria Sebastian', bold=True)
    add_split_row('Test Title: Verify user login with valid credentials', 'Test Execution Date: 06/02/2026', bold=True)
    
    # Row 7 (Description)
    desc_row = add_split_row('Description: To verify that a registered user can successfully log into the system and navigate to the dashboard.', '', bold=True)
    
    # Row 8
    add_merged_row('Pre-Condition: User has valid username and password', bold=True)
    
    # Headers
    hdr_row = table.add_row()
    headers = ["Step", "Test Step", "Test Data", "Expected Result", "Actual Result", "Status(Pass/Fail)"]
    for i, h in enumerate(headers):
        hdr_row.cells[i].text = h
        hdr_row.cells[i].paragraphs[0].runs[0].bold = True
        
    steps_data = [
        (
            "1",
            "Navigate to Home Page URL",
            "http://127.0.0.1:8000",
            "Home Page loads successfully",
            "Home Page is loaded",
            "Pass"
        ),
        (
            "2",
            "Click on Login button",
            "N/A",
            "Redirected to Login page",
            "Login page is displayed",
            "Pass"
        ),
        (
            "3",
            "Enter valid email and password",
            "Email: alan@gmail.com\nPassword: alan@123",
            "Credentials are typed into input fields",
            "Credentials successfully entered",
            "Pass"
        ),
        (
            "4",
            "Click on Submit Login Form button",
            "N/A",
            "User is logged in and redirected to Home Page, Profile icon appears",
            "Redirected to Home page and Profile icon visible",
            "Pass"
        ),
        (
            "5",
            "Click on Profile icon to show dropdown",
            "N/A",
            "Dropdown menu becomes visible",
            "Dropdown menu appears",
            "Pass"
        ),
        (
            "6",
            "Click on Dashboard link from dropdown menu",
            "N/A",
            "Navigating to Dashboard page",
            "Redirected to Dashboard",
            "Pass"
        ),
        (
            "7",
            "Verify current URL is Dashboard",
            "URL contains '/dashboard/'",
            "URL successfully verified and confirmation printed",
            "URL contains '/dashboard/'",
            "Pass"
        )
    ]
    
    for row_data in steps_data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text
            
    # Post-condition
    add_merged_row('Post-Condition: User is logged into the system and dashboard page is displayed.', bold=True)
    
    # Adjust widths
    for cell in table.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in table.columns[1].cells:
        cell.width = Inches(1.5)
    for cell in table.columns[2].cells:
        cell.width = Inches(1.2)
    for cell in table.columns[3].cells:
        cell.width = Inches(1.5)
    for cell in table.columns[4].cells:
        cell.width = Inches(1.5)
    for cell in table.columns[5].cells:
        cell.width = Inches(0.8)
        
    # Bold the specific prefixes if needed
    
    doc.save('Login_Test_Case.docx')
    print("Document successfully created!")

if __name__ == "__main__":
    main()
