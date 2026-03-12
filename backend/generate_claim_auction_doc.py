from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime

def main():
    doc = Document()
    
    # Title
    doc.add_heading('Test Report', level=1)
    
    # Create table with 6 columns
    table = doc.add_table(rows=0, cols=6)
    table.style = 'Table Grid'
    
    def add_merged_row(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
        row = table.add_row()
        cell = row.cells[0]
        for i in range(1, 6):
            cell.merge(row.cells[i])
        cell.text = text
        if bold:
            cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = align
        return cell

    def add_split_row(text1, text2, bold=False):
        row = table.add_row()
        # merge 0,1,2
        cell1 = row.cells[0]
        cell1.merge(row.cells[1])
        cell1.merge(row.cells[2])
        cell1.text = text1
        if bold and cell1.paragraphs[0].runs:
            cell1.paragraphs[0].runs[0].bold = True
        
        # merge 3,4,5
        cell2 = row.cells[3]
        cell2.merge(row.cells[4])
        cell2.merge(row.cells[5])
        cell2.text = text2
        if bold and cell2.paragraphs[0].runs:
            cell2.paragraphs[0].runs[0].bold = True
        return row
    
    # Row 0
    add_merged_row('Test Case 4', bold=True)
    # Row 1
    add_merged_row('Project Name: Bidverse', bold=True)
    # Row 2
    add_merged_row('Claim Won Auction Test Case', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    current_date = datetime.datetime.now().strftime("%d/%m/%Y")
    
    # Row 3
    add_split_row('Test Case ID: Test_4', 'Test Designed By: Jewel Treasa Raphel', bold=True)
    add_split_row('Test Priority(Low/Medium/High): High', f'Test Designed Date: {current_date}', bold=True)
    add_split_row('Module Name: Dashboard / Claiming Module', 'Test Executed By: Sona Maria Sebastian', bold=True)
    add_split_row('Test Title: Verify buyer can claim a won auction', f'Test Execution Date: {current_date}', bold=True)
    
    # Row 7 (Description)
    desc_row = add_split_row('Description: To verify that a logged in buyer can navigate to the won auctions section and successfully click the Claim Now button for an eligible auction.', '', bold=True)
    
    # Row 8
    add_merged_row('Pre-Condition: User is logged in as a buyer and has at least one unclaimed won auction.', bold=True)
    
    # Headers
    hdr_row = table.add_row()
    headers = ["Step", "Test Step", "Test Data", "Expected Result", "Actual Result", "Status(Pass/Fail)"]
    for i, h in enumerate(headers):
        hdr_row.cells[i].text = h
        hdr_row.cells[i].paragraphs[0].runs[0].bold = True
        
    steps_data = [
        (
            "1",
            "Login as buyer",
            "Email: jeweltreasaraphel2028@mca.ajce.in\nPassword: 123jewel",
            "User successfully logs in",
            "Login successful",
            "Pass"
        ),
        (
            "2",
            "Navigate to Dashboard 'Won Auctions' section",
            "URL: http://127.0.0.1:8000/dashboard/?section=won",
            "Won Auctions section loads successfully",
            "Won Auctions loaded",
            "Pass"
        ),
        (
            "3",
            "Locate an unclaimed won auction and click 'Claim Now' button",
            "N/A",
            "Action triggered to claim the item",
            "'Claim Now' button clicked",
            "Pass"
        ),
        (
            "4",
            "Verify redirection to checkout/order page",
            "URL should contain 'checkout' or 'order'",
            "User is redirected to complete the claim process",
            "Redirected successfully",
            "Pass"
        )
    ]
    
    for row_data in steps_data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text
            
    # Post-condition
    add_merged_row('Post-Condition: The buyer is redirected to the checkout page to finalize the claim.', bold=True)
    
    # Adjust widths
    for cell in table.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in table.columns[1].cells:
        cell.width = Inches(1.5)
    for cell in table.columns[2].cells:
        cell.width = Inches(1.2)
    for cell in table.columns[3].cells:
        cell.width = Inches(1.5)
    for cell in table.columns[4].cells:
        cell.width = Inches(1.5)
    for cell in table.columns[5].cells:
        cell.width = Inches(0.8)
        
    doc.save('Claim_Auction_Test_Case.docx')
    print("Claim Auction test document successfully created!")

if __name__ == "__main__":
    main()
